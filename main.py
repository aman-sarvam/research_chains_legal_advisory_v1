from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from langchain.schema.output_parser import StrOutputParser
import requests
from bs4 import BeautifulSoup
from langchain.schema.runnable import RunnablePassthrough, RunnableLambda
# from langchain.utilities import DuckDuckGoSearchAPIWrapper
import json
import os 
from langchain_community.utilities import BingSearchAPIWrapper
from openai import OpenAI
from docx import Document
from docx.shared import Pt
import re
from dotenv import load_dotenv

load_dotenv()
bing_subscription_key = os.getenv("b69720ecb09d4f4cbd77b39c3269c570")
openai_api_key = os.getenv("OPENAI_API_KEY")
langchain_api_key = os.getenv("LANGCHAIN_API_KEY")

client = OpenAI(api_key=openai_api_key)
os.environ["BING_SEARCH_URL"] = "https://api.bing.microsoft.com/v7.0/search"
os.environ["LANGCHAIN_TRACING_V2"] = "true"
os.environ["LANGCHAIN_PROJECT"] = "RESEARCH_CHAINS"


RESULTS_PER_QUESTION = 2

# ddg_search = DuckDuckGoSearchAPIWrapper()
search = BingSearchAPIWrapper()

def web_search(query: str, num_results: int = RESULTS_PER_QUESTION):
    results = search.results(query, num_results)
    return [r["link"] for r in results]


SUMMARY_TEMPLATE = """{text} 
-----------
Using the above text, answer in short the following question: 
> {question}
-----------
if the question cannot be answered using the text, imply summarize the text. Include all factual information, numbers, stats etc if available."""  
SUMMARY_PROMPT = ChatPromptTemplate.from_template(SUMMARY_TEMPLATE)


def scrape_text(url: str):
    try:
        response = requests.get(url)
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, "html.parser")

            page_text = soup.get_text(separator=" ", strip=True)
            
            return page_text
        else:
            return f"Failed to retrieve the webpage: Status code {response.status_code}"
    except Exception as e:
        print(e)
        return f"Failed to retrieve the webpage: {e}"


scrape_and_summarize_chain = RunnablePassthrough.assign(
    summary = RunnablePassthrough.assign(
    text=lambda x: scrape_text(x["url"])[:10000]
) | SUMMARY_PROMPT | ChatOpenAI(model="gpt-3.5-turbo-1106") | StrOutputParser()
) | (lambda x: f"URL: {x['url']}\n\nSUMMARY: {x['summary']}")

web_search_chain = RunnablePassthrough.assign(
    urls = lambda x: web_search(x["question"])
) | (lambda x: [{"question": x["question"], "url": u} for u in x["urls"]]) | scrape_and_summarize_chain.map()


SEARCH_PROMPT = ChatPromptTemplate.from_messages(
    [
        (
            "user",
            "Write 3 google search queries to search online that form an "
            "objective opinion from the following: {question}\n"
            "You must respond with a list of strings in the following format: "
            '["query 1", "query 2", "query 3"].',
        ),
    ]
)

search_question_chain = SEARCH_PROMPT | ChatOpenAI(temperature=0) | StrOutputParser() | json.loads

full_research_chain = search_question_chain | (lambda x: [{"question": q} for q in x]) | web_search_chain.map()

WRITER_SYSTEM_PROMPT = "You are an AI critical thinker research assistant. Your sole purpose is to write well written, critically acclaimed, objective and structured reports on given text."


RESEARCH_REPORT_TEMPLATE = """Information:
--------
{research_summary}
--------
Using the above information, answer the following question or topic: "{question}" in a detailed report -- \
The report should focus on the answer to the question, should be well structured, informative, \
in depth, with facts and numbers if available. Strictly, ,keep the report in a questions answer form. 
You must write the report with markdown syntax.
You MUST determine your own concrete and valid opinion based on the given information. Do NOT deter to general and meaningless conclusions.
Write all used source urls at the end of the report, and make sure to not add duplicated sources, but only one reference for each.
Please do your best, this is very important to my career."""  

prompt = ChatPromptTemplate.from_messages(
    [
        ("system", WRITER_SYSTEM_PROMPT),
        ("user", RESEARCH_REPORT_TEMPLATE),
    ]
)

def collapse_list_of_lists(list_of_lists):
    content = []
    for l in list_of_lists:
        content.append("\n\n".join(l))
    return "\n\n".join(content)

chain = RunnablePassthrough.assign(
    research_summary= full_research_chain | collapse_list_of_lists
) | prompt | ChatOpenAI(model="gpt-3.5-turbo-1106") | StrOutputParser()

# result = chain.invoke(
#     {
#         "question" : "1)What does company boat lifestyle do?"
#     }
# )

# print(result)


def get_chatgpt_response(input_email):
    
    prompt = (
        f""" You are a legal research advisor agent. This is an email with a research task sent from a client: 
        '{input_email}'. 
        
        Go through the email and give a list of questions that will help generate a research report to reply to this email. 
        Only output the questions in an JSON format : 
        
            "questions": [
                "1) What are the current regulations regarding the import and use of SIM cards in vehicles in India?",
                "2) Do I need to partner with a local telecom provider to operate the SIM cards in my cars, and what are the implications of such partnerships?",
                "3) Are there any specific data privacy regulations in India that pertain to data collected through SIM cards in vehicles?"
            ]
        
        If there are questions mentioned in the email then only keep the questions mentioned in the email.
        Only ouput the list of questions and nothing else.
        
        """
        #Breakdown questions to simpler, more specific queries 
    )
    print("Questions generation prompt:", prompt)

    try:
    
        response = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": prompt,
            }
        
        ],
        temperature= 0.4, 
        model="gpt-3.5-turbo",
    )

       
        return response.choices[0].message.content
    except Exception as e:
        print(f"An error occurred: {e}")
        return "Sorry, I couldn't fetch a response. Please check your network or API settings."



# def process_questions(questions):
#     print("QUESTIONS LIST:", questions)
#     results = {}
#     for question in questions:
#         print("QUESTION:", question)
#         result = chain.invoke({"question": question})
#         print("Result:", result)
#         results[question] = result
#     return results


def process_questions(json_data):
    if isinstance(json_data, str):
        data = json.loads(json_data)
    else:
        data = json_data

    questions = data['questions']
    print("QUESTIONS LIST:", questions)
    results = {}
    for question in questions:
        print("QUESTION:", question)
        result = chain.invoke({"question": question})  # Replace with your actual function call
        print("Result:", result)
        results[question] = result
    return results


ex_input_email = """
Test Case: 
I am a car manufacturer in South Korea, and I want to start a business in India. 
My business requires me to put a Sim card in my car and I want to know how the regulations within India affect my business generate seven questions that I must answer to understand this better

Specific questions: 
1) What are the current regulations regarding the import and use of SIM cards in vehicles in India? , 
2) Do I need to partner with a local telecom provider to operate the SIM cards in my cars, and what are the implications of such partnerships?, 
3) Are there any specific data privacy regulations in India that pertain to data collected through SIM cards in vehicles? 

"""

questions = get_chatgpt_response(ex_input_email)
print(questions)
results = process_questions(questions)


for question, result in results.items():
    print(f"Question: {question}\nResult: {result}\n\n")
    
    
# test_questions = [
#  "1) What are the key provisions of the Indian Personal Data Protection Bill that my mobile app needs to comply with?,"
# "2) How does the Information Technology Act impact the operations of my mobile app, especially regarding user data privacy and security?,"
# "3) Are there specific consent requirements I need to adhere to when collecting personal data from users through the app?"
# ]    


# test_questions = [
#     "1) What are the current regulations regarding the import and use of SIM cards in vehicles in India?", 
#     "2) Do I need to partner with a local telecom provider to operate the SIM cards in my cars, and what are the implications of such partnerships?", 
#     "3) Are there any specific data privacy regulations in India that pertain to data collected through SIM cards in vehicles?"]
# results_test = process_questions(test_questions)
# print(results_test)




def markdown_to_word(doc, text):
    # Split the text into lines for easier processing
    lines = text.split('\n')
    for line in lines:
        # Heading 1
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        # Heading 2
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        # Heading 3
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        # Links and normal text
        else:
            # Replace markdown links with just the URL in parentheses
            line = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1 (\2)', line)
            if line.strip():
                paragraph = doc.add_paragraph()
                paragraph.add_run(line)

def save_research_to_docx(data, file_name='Research_Report.docx'):
    # Create a new Document
    doc = Document()
    doc.core_properties.title = "Detailed Research Report"

    # Define the style for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # Process each question and its markdown content
    for question, markdown_content in data.items():
        doc.add_heading(question, level=1)
        markdown_to_word(doc, markdown_content)

    # Save the document
    doc.save(file_name)
    print(f"Document saved as {file_name}")
    
save_research_to_docx(results)