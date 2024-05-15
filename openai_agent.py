#Yet to be modified for legal advisory research 

import os
from dotenv import load_dotenv

from langchain.agents import AgentExecutor, create_react_agent
from langchain_core.prompts import PromptTemplate
from langchain.agents import initialize_agent, Tool
from langchain.agents import AgentType
from langchain_openai import ChatOpenAI
from langchain.prompts import MessagesPlaceholder
from langchain.memory import ConversationSummaryBufferMemory
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.summarize import load_summarize_chain
from langchain.tools import BaseTool
from pydantic import BaseModel, Field
from typing import Type
from bs4 import BeautifulSoup
import requests
import json
from langchain.schema import SystemMessage
from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
import time
from langchain import hub
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

load_dotenv()
serper_api_key = os.getenv("SERPER_API_KEY")
openai_api_key = os.getenv("OPENAI_API_KEY")

def search(query):
    url = "https://google.serper.dev/search"

    payload = json.dumps({
        "q": query
    })

    headers = {
        'X-API-KEY': serper_api_key,
        'Content-Type': 'application/json'
    }

    response = requests.request("POST", url, headers=headers, data=payload)

    print(response.text)

    return response.text


def scrape_website(objective: str, url: str):
    print("Starting Scraping....")

    chrome_options = Options()
    chrome_options.add_argument("--headless")  # Ensure GUI is off
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")

    service = ChromeService() 
    driver = webdriver.Chrome(service=service, options=chrome_options)


    driver.get(url)
    time.sleep(5) 


    page_source = driver.page_source
    driver.quit()

    soup = BeautifulSoup(page_source, "html.parser")
    text = soup.get_text()
    print("CONTENT:.....", text)

    if len(text) > 10000:
        output = summary(objective, text)
        return output
    else:
        return text


def summary(objective, content):
    llm = ChatOpenAI(temperature=0, model="gpt-3.5-turbo-16k-0613")
    print("Creating summary....")

    text_splitter = RecursiveCharacterTextSplitter(
        separators=["\n\n", "\n"], chunk_size=10000, chunk_overlap=500)
    docs = text_splitter.create_documents([content])
    map_prompt = """
    Write a summary of the following text for {objective}:
    "{text}"
    SUMMARY:
    """
    map_prompt_template = PromptTemplate(
        template=map_prompt, input_variables=["text", "objective"])

    summary_chain = load_summarize_chain(
        llm=llm,
        chain_type='map_reduce',
        map_prompt=map_prompt_template,
        combine_prompt=map_prompt_template,
        verbose=True
    )

    output = summary_chain.run(input_documents=docs, objective=objective)

    return output


class ScrapeWebsiteInput(BaseModel):
    """Inputs for scrape_website"""
    objective: str = Field(
        description="The objective & task that users give to the agent")
    url: str = Field(description="The url of the website to be scraped")


class ScrapeWebsiteTool(BaseTool):
    name = "scrape_website"
    description = "useful when you need to get data from a website url, passing both url and objective to the function; DO NOT make up any url, the url should only be from the search results"
    args_schema: Type[BaseModel] = ScrapeWebsiteInput

    def _run(self, objective: str, url: str):
        return scrape_website(objective, url)

    def _arun(self, url: str):
        raise NotImplementedError("error here")



tools = [
    Tool(
        name="Search",
        func=search,
        description="useful for when you need to gain information about the company. Use targeted questions"
    ),
    ScrapeWebsiteTool(),
]

system_message = SystemMessage(
    content="""You are a world class researcher, who can do detailed research on any company and produce facts based results; 
            you do not make things up, you will try as hard as possible to gather facts & data to back up the research
            
            Please make sure you complete the objective above with the following rules:
            1/ You should do enough research to gather as much information as possible about the objective
            2/ If there are url of relevant links & articles, you will scrape it to gather more information
            3/ After scraping & search, you should think "is there any new things i should search & scraping based on the data I collected to increase research quality?" If answer is yes, continue; But don't do this more than 3 iterations
            4/ You should not make things up, you should only write facts & data that you have gathered
            5/ In the final output, You should include all reference data & links to back up your research; You should include all reference data & links to back up your research
            6/ In the final output, You should include all reference data & links to back up your research; You should include all reference data & links to back up your research"""
)

agent_kwargs = {
    "extra_prompt_messages": [MessagesPlaceholder(variable_name="memory")],
    "system_message": system_message,
}

llm = ChatOpenAI(temperature=0, model="gpt-3.5-turbo-16k-0613")
memory = ConversationSummaryBufferMemory(
    memory_key="memory", return_messages=True, llm=llm, max_token_limit=1000)

agent = initialize_agent(
    tools,
    llm,
    agent=AgentType.OPENAI_FUNCTIONS,
    verbose=True,
    agent_kwargs=agent_kwargs,
    memory=memory,
    handle_parsing_errors=True
)


def generate_company_questions(company_name):
    questions = [
        "What sector does {company} operate in?",
        "What is the turnover and date of incorporation of {company}?",
        "Who are the directors of {company}?",
        "What is the main business of {company}?",
        "What are the products or services of {company}?",
        "Where is the registered office of {company}?",
        "Are there any other locations {company} operates in?",
        "How many employees does {company} have?",
        "Who is in the main management of {company}?"
    ]

    customized_questions = [question.format(company=company_name) for question in questions]
    
    return customized_questions


# result = agent({"input": query})

company_name = input("Enter official company name: ")
questions = generate_company_questions(company_name)

answers = {}
for question in questions:
    result = agent({"input": question})
    print("result: ", result)
    answers[question] = result['output']

print("-----------Final answer:", answers)
# print("AgentType.OPENAI_FUNCTIONS:",result['output'])



def save_research_as_docx(data, file_name='Company_Research.docx'):
    # Create a new Document
    doc = Document()
    doc.core_properties.title = "Company Research"

    # Define the style for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # Process each question and answer
    for question, answer in data.items():
        doc.add_paragraph(question, style='Heading 2')
        
        # Clean up any markdown links from the answer
        clean_answer = re.sub(r'\[Link\]\((.*?)\)', r'\1', answer)
        
        # Add the answer as a new paragraph
        paragraph = doc.add_paragraph()
        paragraph.add_run(clean_answer)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Save the document
    doc.save(file_name)
    print(f"Document saved as {file_name}")



save_research_as_docx(answers)



# prompt = hub.pull("hwchase17/react")
# agent = create_react_agent(llm, tools, prompt)

# agent_executor = AgentExecutor(agent=agent, tools=tools, verbose=True)

# result_two = agent_executor.invoke({"input": query})
# print("React Agent:", result_two)


